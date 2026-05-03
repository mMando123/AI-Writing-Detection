"""
file_exporter.py — مصدّر الملفات v4
يحافظ على التنسيق الأصلي والصور عند تعديل النص
"""

import os
from io import BytesIO


class FileExporter:
    """مصدّر الملفات — يدمج النص المعدل مع الحفاظ على التنسيق والصور"""

    def export(self, parsed_doc, modified_paragraphs):
        """
        تصدير المستند مع النص المعدَّل والصور الأصلية.
        
        Args:
            parsed_doc: المستند المُحلَّل (ParsedDocument)
            modified_paragraphs: dict {paragraph_id: new_text}
        
        Returns:
            (bytes, file_extension)
        """
        if parsed_doc.file_type == 'txt':
            return self._export_txt(parsed_doc, modified_paragraphs)
        elif parsed_doc.file_type == 'docx':
            return self._export_docx(parsed_doc, modified_paragraphs)
        elif parsed_doc.file_type == 'pdf':
            return self._export_pdf_as_docx(parsed_doc, modified_paragraphs)
        else:
            raise ValueError(f"نوع غير مدعوم للتصدير: {parsed_doc.file_type}")

    # ───────────────────────────────────────
    #  TXT Export
    # ───────────────────────────────────────
    def _export_txt(self, parsed_doc, modified_paragraphs):
        parts = []
        for elem in parsed_doc.elements:
            if elem.type == 'text':
                new_text = modified_paragraphs.get(elem.id, elem.content)
                parts.append(new_text)
        result = '\n\n'.join(parts)
        return result.encode('utf-8'), 'txt'

    # ───────────────────────────────────────
    #  DOCX Export — المطابقة بالفهرس مع حفظ التنسيق
    # ───────────────────────────────────────
    def _export_docx(self, parsed_doc, modified_paragraphs):
        """
        تصدير ملف Word مع تعديل النص بالفهرس المباشر.
        
        الاستراتيجية:
        1. بناء خريطة: docx_para_index → new_text
        2. فتح الملف الأصلي
        3. تعديل الفقرات مباشرة بالفهرس
        4. الحفاظ على التنسيق والصور
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("مكتبة python-docx غير مثبتة.")

        original_data = parsed_doc.metadata.get('original_stream')
        if not original_data:
            raise ValueError("بيانات المستند الأصلي غير متوفرة")

        # ══════════════════════════════════════════════
        # بناء خريطة الفهرس → النص الجديد
        # ══════════════════════════════════════════════
        index_to_new_text = {}
        for elem in parsed_doc.elements:
            if elem.type == 'text' and elem.id in modified_paragraphs:
                if elem.docx_para_index is not None:
                    index_to_new_text[elem.docx_para_index] = modified_paragraphs[elem.id]

        print(f"[EXPORT] عدد الفقرات المعدلة: {len(index_to_new_text)}")

        if not index_to_new_text:
            print("[EXPORT] تحذير: لا توجد فقرات للتعديل!")

        # ══════════════════════════════════════════════
        # فتح الملف الأصلي وتعديل الفقرات بالفهرس
        # ══════════════════════════════════════════════
        docx_doc = DocxDocument(BytesIO(original_data))

        modified_count = 0
        for para_idx, para in enumerate(docx_doc.paragraphs):
            if para_idx in index_to_new_text:
                new_text = index_to_new_text[para_idx]
                old_text = para.text
                self._replace_paragraph_text_safe(para, new_text)
                modified_count += 1
                print(f"[EXPORT] فقرة {para_idx}: '{old_text[:40]}...' → '{new_text[:40]}...'")

        print(f"[EXPORT] تم تعديل {modified_count} فقرة بنجاح")

        # حفظ
        output = BytesIO()
        docx_doc.save(output)
        output.seek(0)
        return output.read(), 'docx'

    def _replace_paragraph_text_safe(self, paragraph, new_text):
        """
        استبدال نص فقرة بطريقة آمنة تحافظ على الصور والتنسيق.
        
        الاستراتيجية المحسّنة:
        1. فصل الـ runs إلى: نصية / صور
        2. إذا كان هناك run نصي واحد فقط → استبداله مباشرة (يحافظ على تنسيقه)
        3. إذا كان هناك عدة runs نصية → توزيع النص الجديد مع الحفاظ
           على تنسيق الـ run الأول وتفريغ الباقي
        4. إذا كانت كل الـ runs صور → إضافة run نصي جديد
        """
        runs = paragraph.runs
        if not runs:
            # لا توجد runs — إنشاء run جديد
            paragraph.clear()
            paragraph.add_run(new_text)
            return

        # فصل الـ runs النصية عن الصور
        text_runs = []
        image_runs = []
        for run in runs:
            if self._run_has_image(run):
                image_runs.append(run)
            else:
                text_runs.append(run)

        if not text_runs:
            # كل الـ runs تحتوي صور — إضافة run نصي جديد قبل الصور
            new_run = paragraph.add_run(new_text)
            # نقل الـ run الجديد قبل أول عنصر صورة
            first_img_elem = image_runs[0]._element
            paragraph._element.insert(
                list(paragraph._element).index(first_img_elem),
                new_run._element
            )
            return

        # ════ الحالة الرئيسية: تعديل الـ runs النصية ════
        
        # وضع النص الجديد في الـ Run الأول (يحافظ على تنسيقه: خط، حجم، لون...)
        text_runs[0].text = new_text

        # تفريغ بقية الـ Runs النصية (لمنع تكرار النص)
        for run in text_runs[1:]:
            run.text = ""

    def _run_has_image(self, run):
        """التحقق هل الـ Run يحتوي على صورة"""
        xml = run._element.xml
        return 'blip' in xml or 'drawing' in xml or 'pict' in xml

    # ───────────────────────────────────────
    #  PDF → DOCX Export
    # ───────────────────────────────────────
    def _export_pdf_as_docx(self, parsed_doc, modified_paragraphs):
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("مكتبة python-docx غير مثبتة.")

        docx_doc = DocxDocument()
        style = docx_doc.styles['Normal']
        style.font.size = Pt(12)

        for elem in parsed_doc.elements:
            if elem.type == 'text':
                new_text = modified_paragraphs.get(elem.id, elem.content)
                para = docx_doc.add_paragraph(new_text)
                if any('\u0600' <= c <= '\u06FF' for c in new_text):
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif elem.type == 'image':
                if elem.image_data and len(elem.image_data) > 100:
                    try:
                        img_stream = BytesIO(elem.image_data)
                        docx_doc.add_picture(img_stream, width=Inches(5))
                    except Exception:
                        docx_doc.add_paragraph("[صورة]")
                else:
                    docx_doc.add_paragraph("[صورة]")

        output = BytesIO()
        docx_doc.save(output)
        output.seek(0)
        return output.read(), 'docx'

    def get_output_filename(self, original_filename, output_type):
        name = os.path.splitext(original_filename)[0]
        return f"{name}_bypassed.{output_type}"
