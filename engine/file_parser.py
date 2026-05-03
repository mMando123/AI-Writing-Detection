"""
file_parser.py — محرك استخراج النصوص v3
يحافظ على أماكن الصور ويخزّن فهرس الفقرة للتصدير الدقيق
"""

import os
import re
import uuid
from io import BytesIO


class DocumentElement:
    """عنصر في المستند — إما نص أو صورة"""
    def __init__(self, elem_type, content=None, image_data=None, image_ext=None, metadata=None):
        self.id = str(uuid.uuid4())[:8]
        self.type = elem_type
        self.content = content
        self.image_data = image_data
        self.image_ext = image_ext
        self.metadata = metadata or {}
        self.docx_para_index = None  # فهرس الفقرة في الملف الأصلي


class ParsedDocument:
    """مستند مُحلَّل يحتوي على عناصر مرتبة (نص + صور)"""
    def __init__(self, filename, file_type):
        self.filename = filename
        self.file_type = file_type
        self.elements = []
        self.raw_text = ""
        self.paragraphs = []
        self.metadata = {}

    def add_text(self, text, docx_para_index=None):
        if text and text.strip():
            elem = DocumentElement('text', content=text.strip())
            elem.docx_para_index = docx_para_index
            self.elements.append(elem)
            return elem
        return None

    def add_image(self, image_data, image_ext='png', metadata=None):
        elem = DocumentElement('image', image_data=image_data,
                               image_ext=image_ext, metadata=metadata)
        self.elements.append(elem)
        return elem

    def build_text_and_paragraphs(self):
        text_parts = []
        self.paragraphs = []
        for elem in self.elements:
            if elem.type == 'text':
                text_parts.append(elem.content)
                self.paragraphs.append({
                    'id': elem.id,
                    'text': elem.content,
                    'word_count': len(elem.content.split()),
                    'docx_para_index': elem.docx_para_index
                })
        self.raw_text = '\n\n'.join(text_parts)

    def get_image_count(self):
        return sum(1 for e in self.elements if e.type == 'image')


class FileParser:
    """محلل الملفات — يستخرج النص والصور مع الحفاظ على الترتيب"""

    SUPPORTED_TYPES = {'.docx', '.pdf', '.txt'}
    MAX_FILE_SIZE = 20 * 1024 * 1024

    def parse(self, file_path=None, file_stream=None, filename=None):
        if file_path:
            filename = filename or os.path.basename(file_path)
            with open(file_path, 'rb') as f:
                file_data = f.read()
        elif file_stream:
            file_data = file_stream.read()
            if hasattr(file_stream, 'seek'):
                file_stream.seek(0)
        else:
            raise ValueError("يجب تقديم file_path أو file_stream")

        ext = os.path.splitext(filename)[1].lower()
        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"نوع الملف غير مدعوم: {ext}")

        if len(file_data) > self.MAX_FILE_SIZE:
            raise ValueError(f"حجم الملف كبير جداً. الحد الأقصى: {self.MAX_FILE_SIZE // (1024*1024)} ميجابايت")

        if ext == '.txt':
            doc = self._parse_txt(file_data, filename)
        elif ext == '.docx':
            doc = self._parse_docx(file_data, filename)
        elif ext == '.pdf':
            doc = self._parse_pdf(file_data, filename)
        else:
            raise ValueError(f"نوع غير مدعوم: {ext}")

        doc.build_text_and_paragraphs()
        return doc

    def _parse_txt(self, file_data, filename):
        doc = ParsedDocument(filename, 'txt')
        text = None
        for encoding in ['utf-8', 'utf-8-sig', 'cp1256', 'iso-8859-6', 'latin-1']:
            try:
                text = file_data.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if text is None:
            text = file_data.decode('utf-8', errors='replace')

        paragraphs = re.split(r'\n\s*\n', text)
        for i, para in enumerate(paragraphs):
            clean = para.strip()
            if clean:
                doc.add_text(clean, docx_para_index=i)
        return doc

    def _parse_docx(self, file_data, filename):
        """تحليل ملف Word — يخزّن فهرس كل فقرة لضمان التصدير الدقيق"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("مكتبة python-docx غير مثبتة. قم بتشغيل: pip install python-docx")

        doc = ParsedDocument(filename, 'docx')
        doc.metadata['original_stream'] = file_data

        docx_doc = DocxDocument(BytesIO(file_data))

        # استخراج الصور
        image_rels = {}
        for rel_id, rel in docx_doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    image_data = image_part.blob
                    content_type = image_part.content_type
                    ext = content_type.split('/')[-1] if '/' in content_type else 'png'
                    if ext == 'jpeg':
                        ext = 'jpg'
                    image_rels[rel_id] = {'data': image_data, 'ext': ext}
                except Exception:
                    continue

        # المرور على الفقرات بالترتيب — مع حفظ الفهرس!
        images_added = set()
        for para_idx, para in enumerate(docx_doc.paragraphs):
            para_text = para.text.strip()

            # التحقق هل الفقرة تحتوي على صورة
            para_xml = para._element.xml
            has_image = 'blip' in para_xml or 'drawing' in para_xml

            if has_image:
                import re as _re
                rids = _re.findall(r'r:embed="(rId\d+)"', para_xml)
                for rid in rids:
                    if rid in image_rels and rid not in images_added:
                        img = image_rels[rid]
                        doc.add_image(img['data'], img['ext'], {'rel_id': rid})
                        images_added.add(rid)

            # إضافة النص مع فهرس الفقرة
            if para_text:
                doc.add_text(para_text, docx_para_index=para_idx)

        # استخراج النص من الجداول
        for table in docx_doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    doc.add_text(' | '.join(row_texts), docx_para_index=None)

        return doc

    def _parse_pdf(self, file_data, filename):
        try:
            import fitz
        except ImportError:
            raise ImportError("مكتبة PyMuPDF غير مثبتة. قم بتشغيل: pip install PyMuPDF")

        doc = ParsedDocument(filename, 'pdf')
        doc.metadata['original_bytes'] = file_data

        pdf_doc = fitz.open(stream=file_data, filetype="pdf")

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            
            # استخراج النص العادي
            page_text = page.get_text("text")
            has_text = bool(page_text and len(page_text.strip()) > 15)
            
            # استخراج الصور
            image_list = page.get_images(full=True)
            for img_info in image_list:
                xref = img_info[0]
                try:
                    base_image = pdf_doc.extract_image(xref)
                    if base_image and base_image.get("image"):
                        doc.add_image(
                            base_image["image"],
                            base_image.get("ext", "png"),
                            {'page': page_num, 'xref': xref}
                        )
                except Exception:
                    continue

            # ═══ محاولة استخدام OCR (المسح الضوئي) إذا لم يكن هناك نص ═══
            text_to_process = page_text
            if not has_text and len(image_list) > 0:
                try:
                    import pytesseract
                    from PIL import Image
                    print(f"[OCR] جاري مسح الصفحة {page_num} ضوئياً...")
                    # تحويل الصفحة إلى صورة عالية الدقة
                    pix = page.get_pixmap(dpi=150)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    # استخراج النص العربي والإنجليزي
                    ocr_text = pytesseract.image_to_string(img, lang='ara+eng')
                    if ocr_text and len(ocr_text.strip()) > 10:
                        text_to_process = ocr_text
                        print(f"[OCR] نجح المسح في الصفحة {page_num}!")
                except ImportError:
                    print("[OCR] مكتبة pytesseract غير متوفرة.")
                except Exception as e:
                    print(f"[OCR] فشل في الصفحة {page_num}: تأكد من تثبيت برنامج Tesseract-OCR على جهازك. {e}")

            # تقسيم النص إلى فقرات
            if text_to_process and text_to_process.strip():
                import re
                paragraphs = re.split(r'\n\s*\n', text_to_process.strip())
                for para in paragraphs:
                    clean = para.strip().replace('\n', ' ')
                    if clean and len(clean) > 3:
                        doc.add_text(clean)

        pdf_doc.close()
        return doc
